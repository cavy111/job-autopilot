"""
agents/cover_letter.py

Generates a tailored cover letter as a .docx file for a specific job.

What it produces:
  - Professional cover letter in Calvin's voice
  - 3-4 paragraphs: hook, relevant experience, skills fit, closing
  - Saved as a .docx with matching navy blue Arial theme
  - Filename: CoverLetter_Company_Role_timestamp.docx

Modes:
  dry_run=True  — prints prompt only, no API call
  dry_run=False — calls Qwen and generates real .docx
"""

import os
import json
import logging
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

CV_THEME = {
    "name_color":    "1A5276",
    "heading_color": "1A5276",
    "font":          "Arial",
}

# ---------------------------------------------------------------------------
# Prompt builder
# ---------------------------------------------------------------------------

def _build_cover_letter_prompt(cv_profile: dict, job: dict) -> tuple[str, str]:
    """Build system and user prompts for cover letter generation."""

    skills = cv_profile.get("skills", {})
    all_skills = (
        skills.get("languages", [])
        + skills.get("frameworks", [])
        + skills.get("databases", [])
        + skills.get("devops", [])
    )

    exp = cv_profile.get("experience", [])
    exp_text = "\n".join(
        f"- {e.get('title')} at {e.get('company')} ({e.get('period')})"
        for e in exp
    )

    system_prompt = """You are a professional cover letter writer helping a job applicant write a compelling, personalised cover letter.

Write a cover letter with exactly 4 paragraphs:

1. OPENING (2-3 sentences): Express genuine interest in this specific role and company. Mention the role title. Do NOT start with "I am writing to apply..."

2. EXPERIENCE (3-4 sentences): Highlight 1-2 most relevant past experiences that directly relate to this job. Be specific — mention real technologies, outcomes, or responsibilities that match the job.

3. SKILLS FIT (2-3 sentences): Connect specific skills from the candidate's profile to the requirements of this role. Be concrete, not generic.

4. CLOSING (2-3 sentences): Express enthusiasm for contributing to the company, mention availability for interview, and sign off professionally.

Rules:
- Write in first person, professional but warm tone
- NEVER invent experience or skills not in the candidate profile
- Be specific to THIS company and role — not a generic letter
- Keep total length to 250-320 words
- Return ONLY a JSON object with no markdown or explanation

Return this exact JSON:
{
  "subject_line": "Application for [Role Title] — [Candidate Name]",
  "body": "full cover letter text here, paragraphs separated by \\n\\n",
  "word_count": <integer>
}"""

    user_message = f"""CANDIDATE PROFILE:
Name: {cv_profile.get('name')}
Email: {cv_profile.get('email')}
Phone: {cv_profile.get('phone')}
Summary: {cv_profile.get('summary', '')[:400]}
Skills: {', '.join(all_skills)}
Experience:
{exp_text}
Education: {cv_profile.get('education', [{}])[0].get('degree', '')} ({cv_profile.get('education', [{}])[0].get('grade', '')}) — {cv_profile.get('education', [{}])[0].get('institution', '')}

JOB:
Title: {job.get('title')}
Company: {job.get('company')}
Location: {job.get('location')}
Type: {job.get('job_type')}
Description: {(job.get('description') or 'Not provided — write based on job title and company')[:1000]}

Write a tailored cover letter for this specific role."""

    return system_prompt, user_message


# ---------------------------------------------------------------------------
# LLM call
# ---------------------------------------------------------------------------

def _call_llm(system_prompt: str, user_message: str) -> dict:
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai not installed. Run: pip install openai")

    api_key = os.getenv("QWEN_API_KEY")
    if not api_key:
        raise EnvironmentError("QWEN_API_KEY not set in .env")

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
    )

    logger.info("Calling Qwen to generate cover letter...")
    response = client.chat.completions.create(
        model="qwen-turbo",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_message},
        ],
        temperature=0.6,   # higher for natural writing
        max_tokens=1000,
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    return json.loads(raw)


# ---------------------------------------------------------------------------
# .docx generation
# ---------------------------------------------------------------------------

def _build_docx(cv_profile: dict, letter: dict, job: dict, output_path: str):
    """Generate the cover letter as a .docx file."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    navy  = RGBColor(0x1A, 0x52, 0x76)
    black = RGBColor(0x00, 0x00, 0x00)
    font  = CV_THEME["font"]

    def para(text, size=11, bold=False, color=black, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name  = font
        run.font.size  = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        return p

    # ── Header: Name + contact ──
    para(cv_profile.get("name", ""), size=14, bold=True, color=navy,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    contact = f"{cv_profile.get('phone', '')}  ·  {cv_profile.get('email', '')}"
    para(contact, size=10, color=black, space_after=2)

    # ── Divider line ──
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), CV_THEME["heading_color"])
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Date ──
    date_str = datetime.now().strftime("%d %B %Y")
    para(date_str, size=10, space_after=4)

    # ── Recipient ──
    para(f"The Hiring Manager", size=11, space_after=2)
    para(f"{job.get('company', '')}", size=11, space_after=2)
    para(f"{job.get('location', '')}", size=11, space_after=8)

    # ── Subject line ──
    para(letter.get("subject_line", ""), size=11, bold=True, color=navy, space_after=8)

    # ── Body paragraphs ──
    body = letter.get("body", "")
    for paragraph in body.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            para(paragraph, size=11, space_after=8)

    # ── Sign off ──
    para("Yours sincerely,", size=11, space_after=16)
    para(cv_profile.get("name", ""), size=11, bold=True, space_after=2)
    para(cv_profile.get("phone", ""), size=10, space_after=2)
    para(cv_profile.get("email", ""), size=10, space_after=2)

    doc.save(output_path)
    logger.info(f"Cover letter saved to: {output_path}")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_cover_letter(
    cv_profile: dict,
    job: dict,
    output_dir: str = "output/cover_letters",
    dry_run: bool = False,
) -> str:
    """
    Generate a tailored cover letter for a specific job.

    Args:
        cv_profile:  Parsed CV dict from cv_parser.parse_cv()
        job:         Job listing dict from scraper
        output_dir:  Where to save the .docx file
        dry_run:     If True, print prompt only — no API call or file output

    Returns:
        Path to generated .docx, or empty string in dry_run mode.
    """
    system_prompt, user_message = _build_cover_letter_prompt(cv_profile, job)

    if dry_run:
        print("\n" + "=" * 60)
        print("DRY RUN — COVER LETTER GENERATOR")
        print("=" * 60)
        print(f"\nJob: {job.get('title')} @ {job.get('company')}")
        print("\n--- SYSTEM PROMPT ---")
        print(system_prompt)
        print("\n--- USER MESSAGE ---")
        print(user_message)
        print("\n[No API call made — set dry_run=False with QWEN_API_KEY to generate]")
        return ""

    if not os.getenv("QWEN_API_KEY"):
        logger.warning("QWEN_API_KEY not set — falling back to dry run")
        return generate_cover_letter(cv_profile, job, output_dir, dry_run=True)

    letter = _call_llm(system_prompt, user_message)
    logger.info(f"Cover letter generated — ~{letter.get('word_count', '?')} words")

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    company_slug = job.get("company", "company").replace(" ", "_")[:20]
    role_slug    = job.get("title",   "role").replace(" ", "_")[:25]
    timestamp    = datetime.now().strftime("%Y%m%d_%H%M")
    filename     = f"CoverLetter_{company_slug}_{role_slug}_{timestamp}.docx"
    output_path  = str(Path(output_dir) / filename)

    _build_docx(cv_profile, letter, job, output_path)
    return output_path


# ---------------------------------------------------------------------------
# Quick test — python agents/cover_letter.py
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    from pathlib import Path
    import json

    # Reuse the same hardcoded profile from cv_tailor.py
    cv_profile = {
        "name": "Dube Calvin",
        "email": "calvindube.cd@gmail.com",
        "phone": "+263 782 821 968",
        "location": "Zimbabwe",
        "summary": (
            "BSc (Hons) Information Systems graduate with over a year of professional "
            "software development and IT systems experience. Proficient in Python, "
            "JavaScript, PHP, and Java with hands-on Django and React experience."
        ),
        "skills": {
            "languages":  ["Python", "JavaScript", "PHP", "Java", "HTML5", "CSS3"],
            "frameworks": ["Django", "React", "Laravel", "Spring Boot"],
            "databases":  ["SQL", "MySQL", "PostgreSQL", "SQLite"],
            "devops":     ["DigitalOcean", "Git", "REST APIs"],
            "other":      ["Full-stack debugging", "IT systems support", "Technical documentation"],
        },
        "experience": [
            {
                "title": "Web Applications Developer",
                "company": "CNBS Accounting Officers",
                "location": "Pretoria, South Africa",
                "period": "May 2024 – June 2025",
                "bullets": [
                    "Provided ongoing IT systems support for internal web applications.",
                    "Diagnosed and resolved software defects across React/Django stack.",
                    "Maintained production systems on DigitalOcean.",
                ],
            },
            {
                "title": "ICT Facilitator",
                "company": "Fountain Junior School",
                "location": "Zimbabwe",
                "period": "February 2026 – April 2026",
                "bullets": [
                    "Managed classroom computer equipment and troubleshot hardware/software issues.",
                    "Delivered ICT support and training to staff and students.",
                ],
            },
        ],
        "education": [
            {
                "degree": "BSc (Hons) Information Systems",
                "grade": "2.1",
                "institution": "Midlands State University",
                "location": "Gweru, Zimbabwe",
                "period": "2017 – 2022",
            }
        ],
        "certifications": [
            "JPMorgan Chase Software Engineering Job Simulation · Forage · May 2026",
        ],
        "references": [],
        "raw_text": "",
    }

    listings_path = Path("test_output.json")
    if listings_path.exists():
        with open(listings_path) as f:
            jobs = json.load(f)
        job = next(
            (j for j in jobs if "Technology" in j.get("title", "")),
            jobs[0]
        )
    else:
        job = {
            "title": "Software Developer",
            "company": "TechCorp Zimbabwe",
            "location": "Harare",
            "job_type": "Full Time",
            "url": "https://vacancymail.co.zw/jobs/test/",
            "description": "Python/Django developer with React experience needed.",
        }

    print(f"\nTesting cover letter for: {job.get('title')} @ {job.get('company')}")

    if os.getenv("QWEN_API_KEY"):
        output = generate_cover_letter(cv_profile, job, dry_run=False)
        print(f"\nCover letter saved to: {output}")
    else:
        generate_cover_letter(cv_profile, job, dry_run=True)