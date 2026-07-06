"""
agents/cv_tailor.py

Takes Calvin's parsed CV profile and a job listing, and produces a
tailored CV as a .docx file — rewriting the summary and reordering
skills to match the job description.

What it changes:
  - Professional summary: rewritten to mirror the job's language and priorities
  - Skills section: relevant skills surfaced first
  - Experience bullets: most relevant bullets moved to top per role (future)

What it does NOT change:
  - Job titles, companies, dates (never fabricate)
  - Education, references, certifications

Modes:
  dry_run=True  — prints the prompt and a mock response, no API call
  dry_run=False — calls Qwen API and generates real .docx output
"""

import os
import json
import logging
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# CV styling — matches Calvin's existing navy blue theme
CV_THEME = {
    "name_color":    "1A5276",   # navy blue
    "heading_color": "1A5276",
    "font":          "Arial",
    "accent_line":   True,
}

# ---------------------------------------------------------------------------
# Prompt builder
# ---------------------------------------------------------------------------

def _build_tailor_prompt(cv_profile: dict, job: dict) -> tuple[str, str]:
    """
    Build the system and user prompts for CV tailoring.
    Returns (system_prompt, user_message).
    """
    skills = cv_profile.get("skills", {})
    all_skills = (
        skills.get("languages", [])
        + skills.get("frameworks", [])
        + skills.get("databases", [])
        + skills.get("devops", [])
        + skills.get("other", [])
    )

    exp = cv_profile.get("experience", [])
    exp_text = "\n".join(
        f"- {e.get('title')} at {e.get('company')} ({e.get('period')}): "
        + "; ".join(e.get("bullets", [])[:3])
        for e in exp
    )

    system_prompt = """You are a professional CV writer helping a job applicant tailor their CV summary and skills for a specific role.

Your task:
1. Rewrite the candidate's professional summary (3-4 sentences) to:
   - Mirror the language and priorities of the job description
   - Highlight the most relevant skills and experience for THIS specific role
   - Sound natural and authentic, not keyword-stuffed
   - End with enthusiasm for this specific company/role

2. Reorder the candidate's skills list to surface the most relevant ones first for this job.

Rules:
- NEVER invent experience, skills, or qualifications the candidate doesn't have
- Keep the summary truthful — only emphasise what's already in the profile
- The summary should feel personalised, not generic
- Return ONLY a JSON object, no markdown, no explanation

Return this exact JSON structure:
{
  "tailored_summary": "rewritten summary paragraph here",
  "prioritised_skills": {
    "languages": ["ordered list — most relevant first"],
    "frameworks": ["ordered list — most relevant first"],
    "databases": ["ordered list — most relevant first"],
    "devops": ["ordered list — most relevant first"],
    "other": ["ordered list — most relevant first"]
  },
  "tailoring_notes": "1-2 sentences explaining what you emphasised and why"
}"""

    user_message = f"""CANDIDATE PROFILE:
Name: {cv_profile.get('name')}
Current Summary: {cv_profile.get('summary', '')}
Skills: {', '.join(all_skills)}
Experience:
{exp_text}
Education: {cv_profile.get('education', [{}])[0].get('degree', '')} — {cv_profile.get('education', [{}])[0].get('institution', '')}

JOB TO APPLY FOR:
Title: {job.get('title')}
Company: {job.get('company')}
Location: {job.get('location')}
Type: {job.get('job_type')}
Description: {(job.get('description') or 'Not provided — tailor based on job title alone')[:1000]}

Please tailor the CV summary and reorder the skills for this specific role."""

    return system_prompt, user_message


# ---------------------------------------------------------------------------
# LLM call
# ---------------------------------------------------------------------------

def _call_llm(system_prompt: str, user_message: str) -> dict:
    """Send tailoring request to Qwen and return structured response."""
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai not installed. Run: pip install openai")

    api_key = os.getenv("QWEN_API_KEY")
    if not api_key:
        raise EnvironmentError("QWEN_API_KEY not set in .env")

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
    )

    logger.info("Calling Qwen to tailor CV...")
    response = client.chat.completions.create(
        model="qwen3.5-plus",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        temperature=0.4,   # slightly higher for creative rewriting
        max_tokens=1500,
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    return json.loads(raw)


# ---------------------------------------------------------------------------
# .docx generation
# ---------------------------------------------------------------------------

def _build_docx(cv_profile: dict, tailored: dict, job: dict, output_path: str):
    """
    Generate a tailored CV .docx file using python-docx.
    Preserves Calvin's navy blue Arial theme.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    navy = RGBColor(0x1A, 0x52, 0x76)
    black = RGBColor(0x00, 0x00, 0x00)

    def add_heading(text, size=13, color=navy, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.name = CV_THEME["font"]
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return p

    def add_section_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = CV_THEME["font"]
        run.font.size = Pt(10)
        run.font.color.rgb = navy
        # Accent line via bottom border
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), CV_THEME["heading_color"])
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_body(text, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = CV_THEME["font"]
        run.font.size = Pt(size)
        run.font.color.rgb = black
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_bullet(text, size=10):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.name = CV_THEME["font"]
        run.font.size = Pt(size)
        run.font.color.rgb = black
        p.paragraph_format.space_after = Pt(1)
        return p

    # ── Name ──
    p = add_heading(cv_profile.get("name", ""), size=18, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Tagline (from skills) ──
    skills = tailored.get("prioritised_skills", cv_profile.get("skills", {}))
    top_langs = skills.get("languages", [])[:3]
    top_frameworks = skills.get("frameworks", [])[:2]
    tagline = "  ·  ".join(top_langs + top_frameworks)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tagline)
    run.font.name = CV_THEME["font"]
    run.font.size = Pt(10)
    run.font.color.rgb = navy
    run.italic = True

    # ── Contact ──
    contact = f"{cv_profile.get('phone', '')}  ·  {cv_profile.get('email', '')}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(contact)
    run.font.name = CV_THEME["font"]
    run.font.size = Pt(10)
    run.font.color.rgb = black

    doc.add_paragraph()  # spacer

    # ── Professional Summary ──
    add_section_title("Professional Summary")
    add_body(tailored.get("tailored_summary", cv_profile.get("summary", "")))

    doc.add_paragraph()

    # ── Technical Skills ──
    add_section_title("Technical Skills")
    skill_labels = {
        "languages": "Programming Languages",
        "frameworks": "Frameworks",
        "databases": "Databases",
        "devops": "DevOps & Tools",
        "other": "Other Skills",
    }
    for key, label in skill_labels.items():
        items = skills.get(key, [])
        if items:
            p = doc.add_paragraph()
            bold_run = p.add_run(f"{label}:  ")
            bold_run.bold = True
            bold_run.font.name = CV_THEME["font"]
            bold_run.font.size = Pt(10)
            bold_run.font.color.rgb = black
            val_run = p.add_run(", ".join(items))
            val_run.font.name = CV_THEME["font"]
            val_run.font.size = Pt(10)
            val_run.font.color.rgb = black
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ── Work Experience ──
    add_section_title("Work Experience")
    for exp in cv_profile.get("experience", []):
        # Role title | Company
        p = doc.add_paragraph()
        title_run = p.add_run(f"{exp.get('title')}  |  {exp.get('company')}")
        title_run.bold = True
        title_run.font.name = CV_THEME["font"]
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = black

        # Location · Period
        add_body(f"{exp.get('location', '')}  ·  {exp.get('period', '')}")

        for bullet in exp.get("bullets", []):
            add_bullet(bullet)

        doc.add_paragraph()

    # ── Education ──
    add_section_title("Education")
    for edu in cv_profile.get("education", []):
        p = doc.add_paragraph()
        deg = edu.get("degree", "")
        grade = edu.get("grade")
        deg_text = f"{deg} — {grade}" if grade else deg
        run = p.add_run(f"{deg_text}  |  {edu.get('institution', '')}")
        run.bold = True
        run.font.name = CV_THEME["font"]
        run.font.size = Pt(10)
        run.font.color.rgb = black
        add_body(f"{edu.get('location', '')}  ·  {edu.get('period', '')}")

    doc.add_paragraph()

    # ── Certifications ──
    certs = cv_profile.get("certifications", [])
    if certs:
        add_section_title("Certifications & Licences")
        for cert in certs:
            add_bullet(cert)
        doc.add_paragraph()

    # ── References ──
    refs = cv_profile.get("references", [])
    if refs:
        add_section_title("References")
        for ref in refs:
            p = doc.add_paragraph()
            name_run = p.add_run(f"{ref.get('name')}  —  ")
            name_run.bold = True
            name_run.font.name = CV_THEME["font"]
            name_run.font.size = Pt(10)
            name_run.font.color.rgb = black
            role_run = p.add_run(f"{ref.get('role')}  |  {ref.get('contact')}")
            role_run.font.name = CV_THEME["font"]
            role_run.font.size = Pt(10)
            role_run.font.color.rgb = black
            p.paragraph_format.space_after = Pt(2)

    doc.save(output_path)
    logger.info(f"Tailored CV saved to: {output_path}")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def tailor_cv(
    cv_profile: dict,
    job: dict,
    output_dir: str = "output/cvs",
    dry_run: bool = False,
) -> str:
    """
    Generate a tailored CV for a specific job.

    Args:
        cv_profile:  Parsed CV dict from cv_parser.parse_cv()
        job:         Job listing dict from scraper
        output_dir:  Where to save the .docx file
        dry_run:     If True, print prompt and skip API call + docx generation

    Returns:
        Path to the generated .docx file, or empty string in dry_run mode.
    """
    system_prompt, user_message = _build_tailor_prompt(cv_profile, job)

    if dry_run:
        print("\n" + "=" * 60)
        print("DRY RUN — CV TAILOR")
        print("=" * 60)
        print(f"\nJob: {job.get('title')} @ {job.get('company')}")
        print("\n--- SYSTEM PROMPT ---")
        print(system_prompt)
        print("\n--- USER MESSAGE ---")
        print(user_message)
        print("\n[No API call made — set dry_run=False with QWEN_API_KEY to generate real output]")
        return ""

    if not os.getenv("QWEN_API_KEY"):
        logger.warning("QWEN_API_KEY not set — running in dry_run mode instead")
        return tailor_cv(cv_profile, job, output_dir, dry_run=True)

    # Call LLM
    tailored = _call_llm(system_prompt, user_message)
    logger.info(f"Tailoring notes: {tailored.get('tailoring_notes', '')}")

    # Generate output filename
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    company_slug = job.get("company", "company").replace(" ", "_")[:20]
    role_slug    = job.get("title", "role").replace(" ", "_")[:25]
    timestamp    = datetime.now().strftime("%Y%m%d_%H%M")
    filename     = f"CV_{company_slug}_{role_slug}_{timestamp}.docx"
    output_path  = str(Path(output_dir) / filename)

    # Build docx
    _build_docx(cv_profile, tailored, job, output_path)

    return output_path


# ---------------------------------------------------------------------------
# Quick test — python agents/cv_tailor.py
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    # Load a real job from test_output.json
    from pathlib import Path
    import json

    listings_path = Path("test_output.json")
    if listings_path.exists():
        with open(listings_path) as f:
            jobs = json.load(f)
        # Pick the highest-scoring job for testing — Regulatory Tech Developers
        job = next(
            (j for j in jobs if "Technology" in j.get("title", "")),
            jobs[0]
        )
    else:
        job = {
            "title": "Software Developer",
            "company": "TechCorp Zimbabwe",
            "location": "Harare",
            "job_type": "Full Time",
            "url": "https://vacancymail.co.zw/jobs/test/",
            "description": "Looking for a Python/Django developer with React experience.",
        }

    # Use a minimal cv_profile since we don't have Qwen credits yet for cv_parser
    cv_profile = {
        "name": "Dube Calvin",
        "email": "calvindube.cd@gmail.com",
        "phone": "+263 782 821 968",
        "location": "Zimbabwe",
        "summary": (
            "BSc (Hons) Information Systems graduate with over a year of professional "
            "software development and IT systems experience. Skilled in debugging and "
            "troubleshooting software issues across the full stack, managing and maintaining "
            "production systems, and providing technical support to end users. Proficient in "
            "Python, JavaScript, PHP, and Java. Detail-oriented team player who meets deadlines "
            "consistently and approaches technical problems with structured analytical thinking."
        ),
        "skills": {
            "languages":  ["Python", "JavaScript", "PHP", "Java", "HTML5", "CSS3"],
            "frameworks": ["Django", "React", "Laravel", "Spring Boot"],
            "databases":  ["SQL", "MySQL", "PostgreSQL", "SQLite"],
            "devops":     ["DigitalOcean", "Git", "REST APIs"],
            "other":      ["Full-stack debugging", "IT systems support", "Technical documentation"],
        },
        "experience": [
            {
                "title": "Web Applications Developer",
                "company": "CNBS Accounting Officers",
                "location": "Pretoria, South Africa",
                "period": "May 2024 – June 2025",
                "bullets": [
                    "Provided ongoing IT systems support for internal web applications.",
                    "Diagnosed and resolved software defects across React/Django stack.",
                    "Maintained production systems on DigitalOcean.",
                    "Supported end users across accounting, marketing, and development teams.",
                    "Developed and maintained web applications and databases.",
                ],
            },
            {
                "title": "ICT Facilitator",
                "company": "Fountain Junior School",
                "location": "Zimbabwe",
                "period": "February 2026 – April 2026",
                "bullets": [
                    "Managed classroom computer equipment and troubleshot hardware/software issues.",
                    "Delivered ICT support and training to staff and students.",
                ],
            },
            {
                "title": "Laboratory Technician Assistant",
                "company": "Midlands State University",
                "location": "Gweru, Zimbabwe",
                "period": "September 2019 – September 2020",
                "bullets": [
                    "Provided technical support within a university laboratory.",
                    "Maintained accurate records and ensured systems were operational.",
                ],
            },
        ],
        "education": [
            {
                "degree": "BSc (Hons) Information Systems",
                "grade": "2.1",
                "institution": "Midlands State University",
                "location": "Gweru, Zimbabwe",
                "period": "2017 – 2022",
            },
            {
                "degree": "National Certificate in Information Technology",
                "grade": None,
                "institution": "Kwekwe Polytechnic",
                "location": "Kwekwe, Zimbabwe",
                "period": "2016",
            },
        ],
        "certifications": [
            "JPMorgan Chase Software Engineering Job Simulation · Forage · May 2026",
            "Class 4 Driver's Licence",
        ],
        "references": [
            {
                "name": "Mrs Chibwana",
                "role": "Director, Fountain Junior School",
                "contact": "+263 77 321 9259 · Shamisochibwana1975@gmail.com",
            },
            {
                "name": "Mr Giyane",
                "role": "Chairperson, Computer Science Department, Midlands State University",
                "contact": "+263 715 134 137 · giyanem@staff.msu.ac.zw",
            },
            {
                "name": "Mr Phenyo Itumeleng",
                "role": "Senior Developer, CNBS Accounting Officers",
                "contact": "+27 813 280 275 · phenyoitumeleng@gmail.com",
            },
        ],
        "raw_text": "",
    }

    print(f"\nTesting CV tailor for: {job.get('title')} @ {job.get('company')}")

    if os.getenv("QWEN_API_KEY"):
        print("QWEN_API_KEY found — generating real tailored CV...")
        output = tailor_cv(cv_profile, job, dry_run=False)
        print(f"\nTailored CV saved to: {output}")
    else:
        print("No QWEN_API_KEY — running dry run to verify prompt...\n")
        tailor_cv(cv_profile, job, dry_run=True)